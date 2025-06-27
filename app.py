
from flask import Flask, request, send_file
from docx import Document
from tempfile import NamedTemporaryFile

app = Flask(__name__)

@app.route("/generate-docx", methods=["POST"])
def generate_docx():
    data = request.json
    doc = Document()

    doc.add_heading(data.get("job_title", "Job Title"), 0)

    doc.add_heading("Job Purpose", level=1)
    doc.add_paragraph(data.get("job_purpose", ""))

    doc.add_heading("Duties and Responsibilities", level=1)
    doc.add_paragraph(data.get("primary_duties", ""))
    doc.add_paragraph("")  # spacing
    doc.add_paragraph(data.get("additional_duties", ""))

    doc.add_heading("KPIs and Deliverables", level=1)
    doc.add_paragraph(data.get("kpis", ""))

    doc.add_heading("Soft Skills", level=1)
    doc.add_paragraph("See accompanying chart or notes.")

    doc.add_heading("Desired Qualifications / Experiences", level=1)
    doc.add_paragraph(data.get("qualifications", ""))

    doc.add_heading("Working Conditions", level=1)
    doc.add_paragraph(data.get("working_conditions", ""))

    doc.add_heading("Disclaimer", level=1)
    doc.add_paragraph(
        "This job description may not be inclusive of all assigned duties, responsibilities, or aspects of the job described "
        "and may be amended by the employer, at its sole discretion, from time to time to meet evolving business needs. "
        "Employees must be able to satisfactorily perform the essential functions of this position. To meet its obligations "
        "outlined within the provincial Human Rights legislation, the employer is prepared to provide accommodation to the "
        "point of undue hardship to those facing disadvantages related to grounds protected under the law. If reasonable and "
        "justifiable accommodation is required by any employee to perform the essential functions of this position, and/or "
        "receive other benefits and privileges of employment, an accommodation request must be submitted to the Human Resource Department for review."
    )

    # Remove paragraphs containing "Appendix 1" or "Appendix 2"
    for para in doc.paragraphs:
        if para.text.startswith("Appendix 1") or para.text.startswith("Appendix 2"):
            p = para._element
            p.getparent().remove(p)

    temp = NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp.name)
    temp.close()

    return send_file(temp.name, as_attachment=True, download_name=f"{data.get('job_title', 'Job_Description')}.docx")
